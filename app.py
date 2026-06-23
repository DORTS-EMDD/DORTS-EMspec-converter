import streamlit as st
import google.generativeai as genai
import fitz          # PyMuPDF
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import time
import base64
import re
import datetime

# ── 頁面設定 ──────────────────────────────────────────────
st.set_page_config(page_title="未來線中運量機電特別技術規範 - 智慧改寫平台", layout="wide")

for key, val in [
    ("running", False),
    ("result_text", ""),
    ("extracted_old_text", ""),
    ("cf620_total_pages", 0),
    ("cf620_toc", []),
    ("cf620_toc_raw", []),
    ("cf620_detected_offset", 0),
    ("pdf_bytes_cache", None),
    ("chapter_id", ""),
    ("next_chapter_id", ""),
    ("cf620_pdf_name", ""),
    ("kb_cache_key", None),
    ("kb_text_cache", ""),
    ("rewrite_history", []),          # 歷史紀錄：list of {ts, label, text, old_text}
    ("current_old_text_snapshot", ""),# 最新一次改寫時的原文快照
]:
    if key not in st.session_state:
        st.session_state[key] = val

def set_running():
    st.session_state.running = True

# ══════════════════════════════════════════════════════════
# 工具函式
# ══════════════════════════════════════════════════════════

@st.cache_data(show_spinner=False)
def detect_toc_pages(file_bytes, max_scan=20):
    """偵測前 max_scan 頁中目錄頁（章節編號密度高的頁面），回傳 page_no list（0-based）"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    toc_pat = re.compile(r"^\d{1,2}(?:\.\d{1,3}){1,4}")
    toc_pages = []
    for page_no in range(min(len(doc), max_scan)):
        lines = doc[page_no].get_text("text").splitlines()
        hits = sum(1 for ln in lines if toc_pat.match(ln.strip()))
        if hits >= 5:
            toc_pages.append(page_no)
    return toc_pages, len(doc)


@st.cache_data(show_spinner=False)
def render_page_to_png(file_bytes, page_no, dpi=150):
    """將 PDF 指定頁（0-based）render 成 PNG bytes"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc[page_no]
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    return pix.tobytes("png")


@st.cache_data(show_spinner=False)
def build_printed_page_map(file_bytes):
    """
    【方案 A 核心】掃過全部頁面，從頁眉/頁尾抓「實際印刷頁碼」，
    建立 {印刷頁碼: pdf_index(0-based)} 對照表。

    PTS 文件頁眉/頁尾常見：
      - "CF620/ 97"、"DF115/ 3" 這類「代號/頁碼」格式
      - 單獨的阿拉伯數字頁碼
      - 羅馬數字（前言常見，這裡略過不納入對照，避免污染正文頁碼）

    回傳 dict；同一印刷頁碼若重複出現，以「第一次出現」為準。
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    # "CF620/ 97" → 取斜線後的數字當印刷頁碼
    pat_code = re.compile(r'[A-Z]{2}\d{3}\s*/\s*(\d{1,4})')
    # 單獨數字頁碼（整行只有一個數字）
    pat_num  = re.compile(r'^\s*(\d{1,4})\s*$')

    page_map = {}
    for idx in range(len(doc)):
        text = doc[idx].get_text("text")
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        if not lines:
            continue
        # 只看頁面上下緣（頁首前 3 行 + 頁尾後 3 行），降低誤抓內文數字
        candidates = lines[:3] + lines[-3:]
        printed = None
        for ln in candidates:
            m = pat_code.search(ln)
            if m:
                printed = int(m.group(1))
                break
            m = pat_num.match(ln)
            if m:
                # 過濾明顯不合理的純數字（避免抓到年份、條號等）
                val = int(m.group(1))
                if 1 <= val <= len(doc) + 50:
                    printed = val
                    break
        if printed is not None and printed not in page_map:
            page_map[printed] = idx  # 0-based
    return page_map


def make_to_abs(page_map):
    """
    依印刷頁碼對照表，建立「邏輯頁碼 → 絕對 PDF 頁碼(1-based)」轉換函式。
    - 命中：直接查表（精確）
    - 未命中：用最接近的已知錨點做局部線性插補
    """
    known = sorted(page_map)

    def to_abs(logical_p):
        if not known:
            return logical_p
        if logical_p in page_map:
            return page_map[logical_p] + 1  # 命中：精確查表（1-based）
        nearest = min(known, key=lambda k: abs(k - logical_p))
        return page_map[nearest] + 1 + (logical_p - nearest)

    return to_abs


def detect_page_offset(api_key, model_name, file_bytes, toc_page_nos, first_logical_page):
    """
    計算 PDF 物理頁碼與目錄所列邏輯頁碼的偏移量。
    作法：取目錄後第一頁 render 成圖，問 Gemini「這頁的頁碼是幾」，
    offset = (pdf_index + 1) - gemini_回答頁碼
    first_logical_page：目錄中最小的頁碼數字（作為預期值）
    """
    if not toc_page_nos:
        return 0
    first_content_pdf_idx = max(toc_page_nos) + 1   # TOC 最後一頁的下一頁
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    if first_content_pdf_idx >= len(doc):
        return 0

    # 嘗試往後幾頁找到有明確頁碼的頁面
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)

    for probe_idx in range(first_content_pdf_idx, min(first_content_pdf_idx + 5, len(doc))):
        png_bytes = render_page_to_png(file_bytes, probe_idx)
        b64 = base64.b64encode(png_bytes).decode("utf-8")
        prompt_parts = [
            {"inline_data": {"mime_type": "image/png", "data": b64}},
            {"text": (
                "這是一份技術文件的某一頁。"
                "請只回答這頁底部或頁首所顯示的頁碼數字（阿拉伯數字）。"
                "若看不到頁碼，請回答 UNKNOWN。"
                "只輸出一個數字或 UNKNOWN，不要其他文字。"
            )}
        ]
        try:
            resp = model.generate_content([{"parts": prompt_parts}])
            ans = resp.text.strip()
            if ans.isdigit():
                printed = int(ans)
                offset = (probe_idx + 1) - printed   # PDF 1-based - 印刷頁碼
                return offset
        except Exception:
            pass
    return 0


def build_toc_via_vision(api_key, model_name, file_bytes):
    """
    主流程：
      1. 偵測目錄頁
      2. 把每頁 render 成圖送 Gemini Vision → 解析邏輯頁碼目錄
      3. 偵測 PDF 物理頁碼偏移，轉換為絕對 PDF 頁碼
    回傳 (toc, total_pages)
      toc = [[level, "X.X.X 標題", absolute_pdf_page], ...]
    """
    toc_page_nos, total = detect_toc_pages(file_bytes, max_scan=30)

    # 目錄常跨多頁；自動補前後相鄰頁，避免只辨識到中間頁而漏掉首頁/尾頁。
    if not toc_page_nos:
        toc_page_nos = list(range(min(8, total)))
    else:
        start_page = max(0, min(toc_page_nos) - 1)
        end_page = min(total - 1, max(toc_page_nos) + 1)
        toc_page_nos = list(range(start_page, end_page + 1))[:8]

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)

    raw_lines = []
    for page_no in toc_page_nos:
        png_bytes = render_page_to_png(file_bytes, page_no, dpi=120)
        b64 = base64.b64encode(png_bytes).decode("utf-8")
        prompt_parts = [
            {"inline_data": {"mime_type": "image/png", "data": b64}},
            {"text": (
                "這是一份技術規範文件的目錄頁圖片。"
                "請擷取所有目錄條目，每行輸出格式為：\n"
                "章節編號 章節標題 頁碼\n"
                "例如：2.1.12 空調系統與火災警報設備 84\n"
                "只輸出條目清單，不要輸出任何說明或前言。"
                "若圖片中沒有目錄，請輸出：NO_TOC"
            )}
        ]
        try:
            resp = model.generate_content([{"parts": prompt_parts}])
            raw_lines.extend(resp.text.strip().splitlines())
        except Exception:
            pass

    # 解析 Gemini 回傳的文字行
    toc_raw = []
    seen = set()
    pat = re.compile(
        r"^(\d{1,2}(?:\.\d{1,3}){1,5}|第\s*[\d一二三四五六七八九十百]+\s*[章節條篇])"
        r"\s+(.+?)\s+(\d{1,4})\s*$"
    )
    for ln in raw_lines:
        ln = ln.strip()
        if not ln or ln == "NO_TOC":
            continue
        ln_clean = re.sub(r"[.·‥…\s]{4,}", " ", ln).strip()
        m = pat.match(ln_clean)
        if not m:
            continue
        sec_id    = m.group(1).strip()
        title     = m.group(2).strip()
        logical_p = int(m.group(3))
        full_title = f"{sec_id} {title}"
        if full_title in seen:
            continue
        seen.add(full_title)
        dots  = sec_id.count(".")
        level = min(dots + 1, 3)
        toc_raw.append([level, full_title, logical_p])

    # ══════════════════════════════════════════════════════
    # 【方案 A】邏輯頁碼 → 絕對 PDF 頁碼：查表（治本），不再用單一 offset
    # ══════════════════════════════════════════════════════
    page_map = build_printed_page_map(file_bytes)
    to_abs   = make_to_abs(page_map)

    if page_map:
        # 有抓到印刷頁碼 → 逐筆查表 / 插補
        toc = [
            [level, title, max(1, min(to_abs(logical_p), total))]
            for level, title, logical_p in toc_raw
        ]
        # 回傳「命中率」供 UI 顯示信心度
        hit = sum(1 for _, _, lp in toc_raw if lp in page_map)
        coverage = hit / len(toc_raw) if toc_raw else 0.0
        offset = 0  # 已不使用 offset，保留欄位相容
    else:
        # 完全抓不到印刷頁碼 → 退回舊版單一 offset 法（保底）
        first_logical = toc_raw[0][2] if toc_raw else 1
        offset = detect_page_offset(api_key, model_name, file_bytes, toc_page_nos, first_logical)
        toc = [
            [level, title, max(1, logical_p + offset)]
            for level, title, logical_p in toc_raw
        ]
        coverage = 0.0

    # ══════════════════════════════════════════════════════
    # 【方案 C】異常偵測：頁碼超出範圍 / 嚴重非遞增 → 標記需校正
    # ══════════════════════════════════════════════════════
    anomalies = []
    pages_only = [p for _, _, p in toc]
    for i, (lv, title, p) in enumerate(toc):
        if p < 1 or p > total:
            anomalies.append(f"「{title}」頁碼 {p} 超出 1～{total} 範圍")
    # 檢查嚴重逆序（後章節頁碼比前章節小很多）
    for i in range(1, len(pages_only)):
        if pages_only[i] < pages_only[i - 1] - 1:
            anomalies.append(
                f"「{toc[i][1]}」頁碼({pages_only[i]}) 小於前一章節({pages_only[i-1]})"
            )
            break
    # 命中率過低也視為需人工確認
    if toc_raw and coverage < 0.5 and page_map:
        anomalies.append(f"自動比對命中率偏低（{coverage:.0%}），建議人工確認")

    needs_review = len(anomalies) > 0

    return toc, toc_raw, offset, total, page_map, coverage, needs_review, anomalies


@st.cache_data(show_spinner=False)
def build_toc_from_pdf_text(file_bytes, max_scan=25):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total = len(doc)
    raw_lines = []
    toc_pages, _ = detect_toc_pages(file_bytes, max_scan=max_scan)
    scan_pages = toc_pages if toc_pages else list(range(min(total, 8)))
    for page_no in scan_pages:
        raw_lines.extend(doc[page_no].get_text("text").splitlines())

    toc_raw = []
    seen = set()
    toc_pat = re.compile(r"^\s*(\d{1,2}(?:\.\d{1,3}){1,5})\s+(.+?)\s+(\d{1,4})\s*$")
    for line in raw_lines:
        ln = re.sub(r"[.\u2026·•]{2,}", " ", line.strip())
        ln = re.sub(r"\s+", " ", ln)
        m = toc_pat.match(ln)
        if not m:
            continue
        sec_id = m.group(1).strip()
        title = m.group(2).strip(" .\u2026·•")
        logical_p = int(m.group(3))
        if logical_p < 1 or logical_p > total + 100:
            continue
        full_title = f"{sec_id} {title}"
        if full_title in seen:
            continue
        seen.add(full_title)
        level = min(sec_id.count(".") + 1, 3)
        toc_raw.append([level, full_title, logical_p])

    if len(toc_raw) < 3:
        return [], [], 0, total, {}, 0.0, True, ["文字層未解析到足夠目錄項目，建議改用 Gemini Vision。"]

    page_map = build_printed_page_map(file_bytes)
    if page_map:
        to_abs = make_to_abs(page_map)
        toc = [[lv, title, max(1, min(to_abs(logical_p), total))]
               for lv, title, logical_p in toc_raw]
        hit = sum(1 for _, _, lp in toc_raw if lp in page_map)
        coverage = hit / len(toc_raw) if toc_raw else 0.0
    else:
        toc = [[lv, title, max(1, min(logical_p, total))]
               for lv, title, logical_p in toc_raw]
        coverage = 0.0

    anomalies = []
    pages_only = [p for _, _, p in toc]
    for i in range(1, len(pages_only)):
        if pages_only[i] < pages_only[i - 1] - 1:
            anomalies.append(f"「{toc[i][1]}」頁碼({pages_only[i]}) 小於前一章節({pages_only[i-1]})")
            break
    if page_map and coverage < 0.5:
        anomalies.append(f"快速解析頁碼命中率偏低（{coverage:.0%}），建議改用 Gemini Vision。")
    if not page_map:
        anomalies.append("快速解析未找到可對照的印刷頁碼，建議改用 Gemini Vision。")
    needs_review = len(anomalies) > 0
    return toc, toc_raw, 0, total, page_map, coverage, needs_review, anomalies


@st.cache_data(show_spinner=False)
def extract_pages(file_bytes, page_start, page_end):
    """擷取指定頁碼範圍（1-based），保留表格結構，過濾頁首頁尾"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total = len(doc)
    p0 = max(0, page_start - 1)
    p1 = min(total - 1, page_end - 1)
    all_text = ""

    # 頁首頁尾過濾：匹配 "CF620/ 97"、"DF115/ 3" 或純數字頁碼
    hf_pat = re.compile(r'^[A-Z]{2}\d{3}/\s*\d+\s*$|^\d{1,4}\s*$')

    for i in range(p0, p1 + 1):
        page = doc[i]
        all_text += f"\n\n--- 第 {i+1} 頁 ---\n"
        try:
            tables = page.find_tables()
            if tables.tables:
                # 取得所有表格的矩形範圍，避免文字與表格重複輸出
                table_rects = [fitz.Rect(tbl.bbox) for tbl in tables.tables]

                # 先輸出不在表格區域內的文字 blocks（同時過濾頁首頁尾）
                for b in page.get_text("blocks"):
                    if b[6] != 0:          # 略過圖片 block
                        continue
                    b_rect = fitz.Rect(b[:4])
                    if any(b_rect.intersects(tr) for tr in table_rects):
                        continue           # 屬於表格區域，跳過（由下方 markdown 輸出）
                    for ln in b[4].strip().splitlines():
                        if not hf_pat.match(ln.strip()):
                            all_text += ln + "\n"

                # 再輸出表格 Markdown（完整保留，包含最後一頁表格）
                for tbl in tables.tables:
                    df = tbl.to_pandas()
                    all_text += "\n[表格]\n" + df.to_markdown(index=False) + "\n"
            else:
                # 無表格：逐 block 輸出，過濾頁首頁尾
                for b in page.get_text("blocks"):
                    if b[6] != 0:
                        continue
                    for ln in b[4].strip().splitlines():
                        if not hf_pat.match(ln.strip()):
                            all_text += ln + "\n"
        except Exception:
            for b in page.get_text("blocks"):
                all_text += b[4].strip() + "\n"
    return all_text, total


def trim_to_chapter(text, chapter_id, next_chapter_id=None):
    """
    從擷取的多頁文字中，精確裁切出目標章節的內容：
    - 往前裁：找到 chapter_id 開頭的那一行，丟棄之前的內容
    - 往後裁：找到 next_chapter_id 開頭的那一行，丟棄之後的內容
    chapter_id 例如 "2.1.12"，允許後面接空格或中文
    """
    if not chapter_id:
        return text

    lines = text.splitlines()

    def make_pat(sec_id):
        # 逸出點號，允許編號後接空格或中文字
        escaped = re.escape(sec_id)
        return re.compile(r"^\s*" + escaped + r"(\s|\u4e00-\u9fff|$)")

    start_pat = make_pat(chapter_id)
    start_idx = None
    for i, ln in enumerate(lines):
        if start_pat.match(ln):
            start_idx = i
            break

    if start_idx is None:
        # 找不到精確起點就原文回傳
        return text

    end_idx = len(lines)
    if next_chapter_id:
        end_pat = make_pat(next_chapter_id)
        for i in range(start_idx + 1, len(lines)):
            if end_pat.match(lines[i]):
                end_idx = i
                break

    return "\n".join(lines[start_idx:end_idx]).strip()


@st.cache_data(show_spinner=False)
def extract_pdf_with_tables(file_bytes):
    all_text = ""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        try:
            tables = page.find_tables()
            page_text = page.get_text("text")
            if tables.tables:
                for tbl in tables.tables:
                    all_text += "\n[表格]\n" + tbl.to_pandas().to_markdown(index=False) + "\n"
            all_text += page_text + "\n"
        except Exception:
            all_text += page.get_text() + "\n"
    return all_text


def image_to_base64(file_bytes):
    return base64.b64encode(file_bytes).decode("utf-8")


def extract_text_from_image(api_key, model_name, img_bytes, mime_type):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    response = model.generate_content([{
        "parts": [
            {"inline_data": {"mime_type": mime_type, "data": image_to_base64(img_bytes)}},
            {"text": "請完整辨識圖片中所有文字與表格，表格以 Markdown 格式（| 欄 | 欄 |）輸出，保留所有數字、單位與編號。"}
        ]
    }])
    return response.text


def result_to_docx(result_text):
    """
    將 AI 改寫結果（=== 區塊格式）轉換為 Word .docx bytes。
    支援：章節標題、一般段落、Markdown 表格、▸/✕/💡 條列。
    """
    doc = docx.Document()

    # ── 全域字體 ────────────────────────────────────────────
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)

    # ── 標題樣式 ────────────────────────────────────────────
    for level, pt, bold in [(1, 14, True), (2, 12, True)]:
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Arial'
        h.font.size = Pt(pt)
        h.font.bold = bold
        h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # 深藍

    # ── 解析 === 區塊 ────────────────────────────────────────
    sections = result_text.split("===")
    for i in range(1, len(sections) - 1, 2):
        title   = sections[i].strip()
        content = sections[i + 1].strip() if i + 1 < len(sections) else ""

        doc.add_heading(title, level=1)

        lines = content.split("\n")
        j = 0
        while j < len(lines):
            line    = lines[j]
            stripped = line.strip()

            # ── Markdown 表格 ────────────────────────────────
            if stripped.startswith("|"):
                tbl_rows = []
                while j < len(lines) and (
                    lines[j].strip().startswith("|") or
                    (not lines[j].strip() and j + 1 < len(lines) and
                     lines[j + 1].strip().startswith("|"))
                ):
                    if lines[j].strip():
                        tbl_rows.append(lines[j].strip())
                    j += 1

                # 過濾分隔行（| --- | --- |）
                data_rows = [r for r in tbl_rows
                             if not re.match(r"^\|[-:\s|]+\|$", r)]

                if data_rows:
                    num_cols = len(data_rows[0].strip().strip("|").split("|"))
                    tbl = doc.add_table(rows=0, cols=num_cols)
                    tbl.style = 'Table Grid'

                    for row_idx, row_text in enumerate(data_rows):
                        cells = [c.strip() for c in
                                 row_text.strip().strip("|").split("|")]
                        row = tbl.add_row()
                        for col_idx, cell_text in enumerate(cells[:num_cols]):
                            cell = row.cells[col_idx]
                            cell.text = cell_text
                            if row_idx == 0:   # 表頭加粗
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.bold = True

                    doc.add_paragraph()   # 表格後空行
                continue   # j 已在內層迴圈推進

            # ── 條列符號行（▸ ✕ 💡）────────────────────────
            elif stripped and stripped[0] in ("▸", "✕", "💡"):
                p = doc.add_paragraph(stripped, style='Normal')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after  = Pt(4)

            # ── 小標（以數字+點+空白開頭，如「1. 」「A. 」）──
            elif re.match(r"^\d{1,2}[\.、]\s|^[A-Z]\.\s", stripped) and stripped:
                p = doc.add_paragraph(stripped, style='Normal')
                p.paragraph_format.left_indent = Inches(0.2)

            # ── 一般文字段落 ─────────────────────────────────
            elif stripped:
                doc.add_paragraph(stripped, style='Normal')

            # ── 空行（視為段落間距，不重複加）────────────────
            else:
                pass   # spacing_after 已由段落格式處理

            j += 1

        # 章節之間加分隔線（水平段落框線）
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        from docx.oxml.ns import qn
        from docx.oxml    import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'BBBBBB')
        pBdr.append(bottom)
        pPr.append(pBdr)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
# ══════════════════════════════════════════════════════════
# 標題 + 全域樣式（限寬置中，避免超寬螢幕字太散）
# ══════════════════════════════════════════════════════════
st.markdown(
    """
    <style>
      .block-container {max-width: 1180px; padding-top: 1.2rem; padding-bottom: 2rem;}
      div[data-testid="stExpander"] {border-radius: 10px;}
      .stTabs [data-baseweb="tab-list"] {gap: 4px;}
      .stTabs [data-baseweb="tab"] {height: 46px; font-size: 1.02rem; font-weight: 600;}
      .pts-hint {background:#f6f8fa;border-left:4px solid #4c8bf5;padding:10px 14px;
                 border-radius:6px;font-size:0.9rem;color:#444;margin-bottom:8px;}
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("🚆 未來線中運量機電特別技術規範 - 智慧改寫平台")
st.caption("運用 Gemini AI，將既有特別技術規範改寫為未來線中運量系統版本 ｜ 三步驟：輸入 → 改寫 → 精修")

# ══════════════════════════════════════════════════════════
# 左側欄：只保留「系統設定 + 精進文件」，提示語移至主畫面
# ══════════════════════════════════════════════════════════
with st.sidebar:
    st.header("⚙️ 系統設定")
    api_key = st.text_input(
        "🔑 Gemini API Key", type="password",
        help="前往 https://aistudio.google.com/app/apikey 取得",
    )
    selected_model = st.selectbox(
        "🤖 選擇模型",
        ["gemini-3.1-flash-lite", "gemini-3.5-flash"],
        index=0,
        help=(
            "gemini-3.1-flash-lite：Gemini 3 最新輕量版，速度快、省配額，適合大量章節改寫。\n"
            "gemini-3.5-flash：接近 Pro 等級智能，細節與表格保留更完整，建議用於關鍵章節。"
        ),
    )
    if api_key:
        st.success("✅ 已輸入 API Key", icon="🔑")
    else:
        st.info("請先輸入 API Key 以啟用辨識與改寫", icon="ℹ️")

    st.markdown("---")
    st.header("📂 精進文件（選填）")
    st.caption("上傳未來線最新測試規定（PDF / Word），AI 改寫時將優先採用其中參數")
    uploaded_files = st.file_uploader(
        "選擇檔案（可多選）",
        accept_multiple_files=True, type=["pdf", "docx"], key="kb_uploader",
    )

# ── 精進文件擷取（邏輯不變，只是搬到 sidebar 區塊外處理） ──
kb_text = ""
if uploaded_files:
    kb_cache_key = tuple((f.name, f.size) for f in uploaded_files)
    if st.session_state.get("kb_cache_key") != kb_cache_key:
        _kb = ""
        for f in uploaded_files:
            try:
                raw = f.read()
                if f.name.lower().endswith(".pdf"):
                    _kb += extract_pdf_with_tables(raw)
                elif f.name.lower().endswith(".docx"):
                    wd = docx.Document(io.BytesIO(raw))
                    for para in wd.paragraphs:
                        _kb += para.text + "\n"
            except Exception as e:
                st.sidebar.warning(f"⚠️ 無法讀取 {f.name}：{e}")
        st.session_state["kb_text_cache"] = _kb
        st.session_state["kb_cache_key"] = kb_cache_key
        st.sidebar.success(f"✅ 成功載入 {len(uploaded_files)} 份精進文件！")
    kb_text = st.session_state.get("kb_text_cache", "")

with st.sidebar:
    if kb_text:
        st.caption(f"📑 目前精進文件約 {len(kb_text):,} 字")

    st.markdown("---")
    with st.expander("🔒 術語保護清單（選填）", expanded=False):
        st.caption("每行一個專業術語，AI 改寫時將原文照錄，不得更動")
        protected_terms_input = st.text_area(
            "受保護術語",
            height=120,
            placeholder="例如：\n750V DC\nCBTC\nATO自動駕駛\nGoA4",
            label_visibility="collapsed",
            key="protected_terms_area",
        )

# ══════════════════════════════════════════════════════════
# 主流程：三個分頁（① 輸入 → ② 改寫結果 → ③ 精修）
# ══════════════════════════════════════════════════════════
tab_input, tab_result, tab_refine = st.tabs([
    "　① 輸入待改寫條文　",
    "　② 改寫結果　",
    "　③ 二次精修　",
])

import streamlit.components.v1 as _stc


def _result_tab_script():
    return """<script>
    (function(){
        const targetText = "② 改寫結果";
        let attempts = 0;

        function getDocs(){
            const docs = [];
            try { docs.push(window.parent.document); } catch(e) {}
            try { docs.push(window.top.document); } catch(e) {}
            try { docs.push(document); } catch(e) {}
            return docs;
        }

        function activate(tab){
            tab.scrollIntoView({block: "center", inline: "center"});
            tab.dispatchEvent(new MouseEvent("mousedown", {bubbles: true, cancelable: true, view: window}));
            tab.dispatchEvent(new MouseEvent("mouseup", {bubbles: true, cancelable: true, view: window}));
            tab.click();
        }

        function trySwitch(){
            attempts += 1;
            for (const d of getDocs()){
                const tabs = Array.from(d.querySelectorAll('[data-baseweb="tab"], button[role="tab"], [role="tab"]'));
                const matched = tabs.find(t => (t.innerText || t.textContent || "").replace(/\\s+/g, " ").includes(targetText));
                if (matched) {
                    activate(matched);
                    return;
                }
                if (tabs.length >= 2) {
                    activate(tabs[1]);
                    return;
                }
            }
            if (attempts < 30) setTimeout(trySwitch, 200);
        }

        setTimeout(trySwitch, 50);
    })();
    </script>"""


def emit_result_tab_switch():
    _stc.html(_result_tab_script(), height=0)


def switch_to_result_page(message="✅ 已切換至「② 改寫結果」", icon="🎉"):
    st.session_state["_switch_to_result"] = {"message": message, "icon": icon}


def generate_with_progress(prompt, model_name, api_key, progress_bar, progress_label):
    retries, success, full_text = 3, False, ""
    while not success and retries > 0:
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt, stream=True)
            full_text = ""
            chunk_count = 0
            progress_bar.progress(0.05, text=f"⏳ {progress_label}中，正在等待 AI 回應…")
            for chunk in response:
                try:
                    full_text += chunk.text
                    chunk_count += 1
                    if chunk_count % 6 == 0:
                        progress_bar.progress(
                            min(chunk_count / 80, 0.95),
                            text=f"⏳ {progress_label}中… 已產出 {len(full_text):,} 字元"
                        )
                except Exception:
                    pass
            progress_bar.progress(1.0, text=f"✅ {progress_label}完成！")
            full_text = re.sub(r'\n{3,}', '\n\n', full_text)
            success = True
        except Exception as e:
            err = str(e)
            if "429" in err or "quota" in err.lower():
                retries -= 1
                if retries > 0:
                    st.warning(f"⚠️ 配額繁忙，15 秒後重試（剩餘 {retries} 次）...")
                    time.sleep(15)
            elif "API_KEY_INVALID" in err or "api key" in err.lower():
                st.error("❌ API Key 無效。")
                break
            else:
                st.error(f"❌ 錯誤：{err}")
                break
    if not success and retries == 0:
        st.error("❌ 已超過重試次數，請稍後手動重試。")
    return success, full_text


# ── 改寫/還原完成後自動跳至 ② 分頁（st.tabs 無原生 API，故用前端點擊）──
_switch_payload = st.session_state.pop("_switch_to_result", None)
if _switch_payload:
    if isinstance(_switch_payload, dict):
        st.toast(_switch_payload.get("message", "✅ 已切換至「② 改寫結果」"), icon=_switch_payload.get("icon", "🎉"))
    else:
        st.toast("✅ 已切換至「② 改寫結果」", icon="🎉")
    emit_result_tab_switch()

# ══════════════════════════════════════════════════════════
# ① 輸入分頁
# ══════════════════════════════════════════════════════════
with tab_input:
    in_main, in_side = st.columns([3, 2], gap="large")

    # ── 左大欄：條文來源（PDF / 貼上 / 截圖） ──
    with in_main:
        st.subheader("📥 條文來源")
        src_pdf, src_paste, src_img = st.tabs([
            "📄 上傳 PDF（推薦）", "✍️ 直接貼上", "🖼️ 截圖辨識",
        ])

        # ── 來源 1：大型 PDF 選頁 ──────────────────────────
        with src_pdf:
            st.markdown(
                '<div class="pts-hint">💡 流程：上傳 PDF → 自動辨識目錄 → 選章節 → 擷取內容</div>',
                unsafe_allow_html=True,
            )
            cf620_pdf = st.file_uploader("上傳 PTS 完整 PDF", type=["pdf"], key="cf620_big_pdf")

            if cf620_pdf:
                if cf620_pdf.name != st.session_state.cf620_pdf_name:
                    pdf_bytes = cf620_pdf.read()
                    st.session_state.pdf_bytes_cache = pdf_bytes
                    st.session_state.cf620_pdf_name = cf620_pdf.name
                    with st.spinner("📖 使用 Gemini Vision 辨識目錄並建立頁碼對照表（最多 8 頁），請稍候…"):
                        if not api_key:
                            st.error("⚠️ 請先在左側輸入 API Key 才能辨識目錄！")
                            toc, toc_raw, detected_offset, total_pages = [], [], 0, 1
                            page_map, coverage, needs_review, anomalies = {}, 0.0, False, []
                        else:
                            (toc, toc_raw, detected_offset, total_pages,
                             page_map, coverage, needs_review, anomalies) = build_toc_via_vision(
                                api_key, selected_model, pdf_bytes)
                        st.session_state.cf620_total_pages     = total_pages
                        st.session_state.cf620_toc             = toc
                        st.session_state.cf620_toc_raw         = toc_raw
                        st.session_state.cf620_detected_offset = detected_offset
                        st.session_state.cf620_page_map        = page_map
                        st.session_state.cf620_coverage        = coverage
                        st.session_state.cf620_needs_review    = needs_review
                        st.session_state.cf620_anomalies       = anomalies
                else:
                    pdf_bytes       = st.session_state.pdf_bytes_cache
                    toc             = st.session_state.cf620_toc
                    toc_raw         = st.session_state.cf620_toc_raw
                    detected_offset = st.session_state.cf620_detected_offset
                    total_pages     = st.session_state.cf620_total_pages
                    page_map        = st.session_state.get("cf620_page_map", {})
                    coverage        = st.session_state.get("cf620_coverage", 0.0)
                    needs_review    = st.session_state.get("cf620_needs_review", False)
                    anomalies       = st.session_state.get("cf620_anomalies", [])

                if st.button("🔁 使用 Gemini Vision 重新辨識目錄", disabled=not api_key, use_container_width=True):
                    if not api_key:
                        st.error("⚠️ 請先在左側輸入 API Key。")
                    else:
                        with st.spinner("📖 使用 Gemini Vision 重新辨識目錄中（最多 8 頁）…"):
                            (toc, toc_raw, detected_offset, total_pages,
                             page_map, coverage, needs_review, anomalies) = build_toc_via_vision(
                                api_key, selected_model, pdf_bytes)
                            st.session_state.cf620_total_pages     = total_pages
                            st.session_state.cf620_toc             = toc
                            st.session_state.cf620_toc_raw         = toc_raw
                            st.session_state.cf620_detected_offset = detected_offset
                            st.session_state.cf620_page_map        = page_map
                            st.session_state.cf620_coverage        = coverage
                            st.session_state.cf620_needs_review    = needs_review
                            st.session_state.cf620_anomalies       = anomalies
                            st.success("✅ Gemini Vision 目錄重新辨識完成")

                # ── 載入摘要 + 頁碼信心度 ──
                c1, c2, c3 = st.columns(3)
                c1.metric("總頁數", total_pages)
                c2.metric("偵測章節", len(toc))
                if page_map:
                    c3.metric("頁碼比對命中", f"{coverage:.0%}")
                else:
                    c3.metric("頁碼比對命中", "—")

                # ══ 方案 C：異常自動提示，命中良好則顯示綠燈 ══
                if needs_review:
                    st.warning(
                        "⚠️ 偵測到頁碼可能有偏差，請展開下方校正面板確認：\n\n- "
                        + "\n- ".join(anomalies[:5])
                    )
                elif page_map and coverage >= 0.8:
                    st.success(f"✅ 頁碼已自動對照完成（命中率 {coverage:.0%}），通常無需手動校正")
                elif not page_map:
                    st.info("ℹ️ 此 PDF 頁眉/頁尾未偵測到可辨識的印刷頁碼，已採用估算偏移，如有偏差請於下方校正")

                # ── 頁碼校正面板 ──
                if toc_raw:
                    with st.expander("⚙️ 頁碼校正（選用）", expanded=False):
                        st.caption("選一個你確定的章節，輸入它在 PDF 中實際的頁碼，程式會自動修正所有章節頁碼")
                        ref_options = [f"第 {item[2]} 頁 | {item[1]}" for item in toc]
                        cr1, cr2 = st.columns([2, 1])
                        with cr1:
                            selected_ref = st.selectbox(
                                "選擇參考章節", ["— 不校正 —"] + ref_options, key="ref_select")
                        with cr2:
                            if selected_ref != "— 不校正 —":
                                actual_page = st.number_input(
                                    "實際頁碼", min_value=1, max_value=max(1, total_pages),
                                    value=int(selected_ref.split()[1]), key="actual_page")
                            else:
                                actual_page = None
                        if selected_ref != "— 不校正 —" and actual_page is not None:
                            ref_idx = ref_options.index(selected_ref)
                            correct_offset = actual_page - toc_raw[ref_idx][2]
                            st.success(f"✅ 校正完成！整體偏移 {correct_offset:+d} 頁")
                            toc = [[lv, t, max(1, lp + correct_offset)] for lv, t, lp in toc_raw]

                # ── 章節選擇 ──
                if toc:
                    toc_options = [f"第{item[2]}頁｜{'　' * (item[0]-1)}{item[1]}" for item in toc]
                    selected_toc = st.selectbox(
                        f"選擇章節（共 {len(toc)} 筆，可輸入關鍵字搜尋）",
                        ["— 請選擇 —"] + toc_options, key="toc_select")
                    with st.expander("需要跨多個章節？", expanded=False):
                        multi_mode = st.checkbox("啟用跨章節選取", key="multi_mode")

                    if not multi_mode:
                        if selected_toc != "— 請選擇 —":
                            chosen_idx = toc_options.index(selected_toc)
                            auto_start = toc[chosen_idx][2]
                            auto_end   = total_pages
                            _next_id   = ""
                            for j in range(chosen_idx + 1, len(toc)):
                                if toc[j][0] <= toc[chosen_idx][0]:
                                    auto_end = toc[j][2] - 1
                                    _next_id = toc[j][1].split()[0]
                                    break
                            auto_end = max(auto_end, auto_start)
                            auto_end = min(auto_end, auto_start + 29)
                            st.session_state.chapter_id      = toc[chosen_idx][1].split()[0]
                            st.session_state.next_chapter_id = _next_id
                        else:
                            auto_start, auto_end = 1, min(10, total_pages)
                            st.session_state.chapter_id = ""
                            st.session_state.next_chapter_id = ""
                    else:
                        with st.expander("跨章節範圍設定", expanded=True):
                            cs, ce = st.columns(2)
                            with cs:
                                start_sel = st.selectbox(
                                    f"起始章節（共 {len(toc)} 筆）",
                                    ["— 請選擇 —"] + toc_options, key="range_start")
                            with ce:
                                end_sel = st.selectbox(
                                    "結束章節", ["— 請選擇 —"] + toc_options, key="range_end")
                            if start_sel != "— 請選擇 —" and end_sel != "— 請選擇 —":
                                si = toc_options.index(start_sel)
                                ei = toc_options.index(end_sel)
                                if ei < si:
                                    st.error("❌ 結束章節不能在起始章節之前，請重新選擇。")
                                    auto_start, auto_end = 1, min(10, total_pages)
                                    st.session_state.chapter_id = ""
                                    st.session_state.next_chapter_id = ""
                                else:
                                    auto_start = toc[si][2]
                                    auto_end   = total_pages
                                    _next_id   = ""
                                    for j in range(ei + 1, len(toc)):
                                        if toc[j][0] <= toc[ei][0]:
                                            auto_end = toc[j][2] - 1
                                            _next_id = toc[j][1].split()[0]
                                            break
                                    auto_end = max(auto_end, auto_start)
                                    page_span = auto_end - auto_start + 1
                                    if page_span > 30:
                                        st.info(
                                            f"ℹ️ 已選 **{toc[si][1]}** → **{toc[ei][1]}**，"
                                            f"共約 **{page_span}** 頁。頁數較多時改寫時間較長。")
                                    st.session_state.chapter_id      = toc[si][1].split()[0]
                                    st.session_state.next_chapter_id = _next_id
                            else:
                                auto_start, auto_end = 1, min(10, total_pages)
                                st.session_state.chapter_id = ""
                                st.session_state.next_chapter_id = ""
                else:
                    st.warning("⚠️ 未偵測到章節標題，請手動輸入頁碼。")
                    auto_start, auto_end = 1, min(10, total_pages)

                # ── 頁碼範圍提示 + 手動微調 ──
                page_count_preview = min(auto_end, total_pages) - auto_start + 1
                if page_count_preview > 30:
                    st.warning(f"⚠️ 預計擷取 {page_count_preview} 頁，建議不超過 30 頁。")
                else:
                    st.success(f"✅ 即將擷取第 {auto_start} ～ {min(auto_end, total_pages)} 頁（共 {page_count_preview} 頁）")

                page_start = max(1, min(auto_start, total_pages)) if total_pages > 0 else 1
                page_end   = max(1, min(auto_end,   total_pages)) if total_pages > 0 else 1

                with st.expander("🔧 手動調整頁碼範圍（選用）", expanded=False):
                    cs2, ce2 = st.columns(2)
                    _safe_start = max(1, min(auto_start, total_pages)) if total_pages > 0 else 1
                    _safe_end   = max(1, min(auto_end,   total_pages)) if total_pages > 0 else 1
                    with cs2:
                        page_start = st.number_input("起始頁", 1, max(1, total_pages), _safe_start)
                    with ce2:
                        page_end = st.number_input("結束頁", 1, max(1, total_pages), _safe_end)
                    if page_end - page_start + 1 > 30:
                        st.warning(f"⚠️ 選取 {page_end - page_start + 1} 頁，建議不超過 30 頁。")

                if st.button("📥 擷取選定頁面內容", use_container_width=True, type="primary"):
                    with st.spinner(f"擷取第 {page_start}～{page_end} 頁..."):
                        try:
                            extracted, _ = extract_pages(pdf_bytes, page_start, page_end)
                            ch_id   = st.session_state.get("chapter_id", "")
                            next_id = st.session_state.get("next_chapter_id", "")
                            if ch_id:
                                extracted = trim_to_chapter(extracted, ch_id, next_id or None)
                            st.session_state.extracted_old_text = extracted
                            st.success(f"✅ 擷取完成！共 {len(extracted)} 字元，可前往「② 改寫結果」執行改寫")
                            with st.expander("預覽擷取內容", expanded=False):
                                st.text(extracted[:3000] + ("..." if len(extracted) > 3000 else ""))
                        except Exception as e:
                            st.error(f"❌ 擷取失敗：{e}")

        # ── 來源 2：手動貼上 ──
        with src_paste:
            _demo_col, _ = st.columns([1, 3])
            with _demo_col:
                if st.button("📋 載入範例條文", help="委員 Demo 時可直接使用，無需輸入原始條文"):
                    st.session_state.extracted_old_text = (
                        "3.1.1 電力供應\n本系統採用 1500V DC 第三軌供電方式，"
                        "變電站間距不超過 2.0 km。\n每節車廂最大用電量為 180 kW。\n\n"
                        "3.1.2 牽引系統\n牽引馬達採用三相感應馬達，額定功率 150 kW/軸，"
                        "最高速度 80 km/h，常用減速度 1.0 m/s²，緊急制動減速度 1.3 m/s²。\n\n"
                        "3.1.3 空調系統\n車廂冷氣量不得低於 12 RT/節，"
                        "設計溫度：夏季室內 26°C，外氣溫度 35°C。"
                    )
                    st.rerun()
            pasted_text = st.text_area(
                "請複製舊規範條文貼在此處：", height=320,
                value=st.session_state.get("extracted_old_text", "") if not st.session_state.get("pdf_bytes_cache") else "",
                placeholder="貼上純文字條文，若有表格無法複製請改用 PDF 或截圖分頁...")

        # ── 來源 3：截圖辨識 ──
        with src_img:
            st.markdown(
                '<div class="pts-hint">📷 截圖後上傳，由 Gemini 辨識圖片中的文字與表格</div>',
                unsafe_allow_html=True)
            cf620_img = st.file_uploader("上傳截圖（PNG / JPG）",
                type=["png", "jpg", "jpeg"], key="cf620_img")
            if cf620_img:
                st.image(cf620_img, use_container_width=True)
                _ic1, _ic2 = st.columns(2)
                with _ic1:
                    if st.button("🔍 執行圖片辨識", use_container_width=True):
                        if not api_key:
                            st.error("⚠️ 請先輸入 API Key！")
                        else:
                            with st.spinner("Gemini 辨識中..."):
                                try:
                                    mime = "image/png" if cf620_img.name.lower().endswith(".png") else "image/jpeg"
                                    result = extract_text_from_image(api_key, selected_model, cf620_img.read(), mime)
                                    st.session_state.extracted_old_text = result
                                    st.success("✅ 辨識完成！可前往「② 改寫結果」執行改寫")
                                    with st.expander("預覽辨識結果", expanded=True):
                                        st.markdown(result)
                                except Exception as e:
                                    st.error(f"❌ 辨識失敗：{e}")
                with _ic2:
                    if st.button("🗑️ 清除辨識結果", use_container_width=True):
                        st.session_state.extracted_old_text = ""
                        st.rerun()

    # ── 右窄欄：改寫方向提示語（從 sidebar 搬來，避免 sidebar 太擠） ──
    with in_side:
        st.subheader("✏️ 改寫方向提示語")
        st.caption("可填改寫方向、特殊要求或精進文字（選填），AI 改寫時一併參考")
        user_hint = st.text_area(
            "提示語", height=260,
            placeholder="例如：\n・電力系統改用 750V DC\n・保留所有測試步驟數值\n・或貼上精進文件全文（約 1000 字內最佳）…",
            label_visibility="collapsed", key="user_hint_area")

        # 來源整合狀態
        st.markdown("##### 📌 目前輸入狀態")
        if st.session_state.extracted_old_text:
            st.success(f"已就緒：PDF / 圖片擷取結果（{len(st.session_state.extracted_old_text)} 字元）")
        else:
            st.info("尚未擷取內容；若用「直接貼上」則於貼上後即可改寫")

    # 整合輸入來源
    if st.session_state.extracted_old_text:
        old_text = st.session_state.extracted_old_text
    else:
        old_text = pasted_text if "pasted_text" in dir() else ""

    st.markdown("---")
    bc1, bc2 = st.columns([3, 1])
    with bc1:
        run_btn = st.button("🚀 執行 AI 改寫（結果顯示於「② 改寫結果」）",
            disabled=st.session_state.running, use_container_width=True, type="primary")
    with bc2:
        clear_btn = st.button("🗑️ 清除結果", use_container_width=True)
    rewrite_status_box = None
    rewrite_progress_bar = None
    _rewrite_can_start = bool(run_btn and api_key and old_text and old_text.strip())
    if run_btn:
        if _rewrite_can_start:
            emit_result_tab_switch()
        else:
            rewrite_status_box = st.empty()
            rewrite_progress_bar = st.progress(0, text="⏳ AI 改寫準備中，請稍候…")

# ══════════════════════════════════════════════════════════
# ② 改寫結果分頁
# ══════════════════════════════════════════════════════════
with tab_result:
    st.subheader("✨ 智慧改寫草稿")
    _pending_refine_prompt = st.session_state.get("_pending_refine_prompt")
    _result_task_active = _rewrite_can_start or bool(_pending_refine_prompt)
    result_progress_bar = None
    if _result_task_active:
        task_label = "二次精修" if _pending_refine_prompt else "AI 改寫"
        result_progress_bar = st.progress(0, text=f"⏳ {task_label}準備中，請稍候…")
        st.caption("任務完成後會在此頁顯示最新結果。")

    if clear_btn:
        st.session_state.result_text = ""
        st.session_state.extracted_old_text = ""
        st.session_state.rewrite_history = []
        st.session_state.current_old_text_snapshot = ""
        st.rerun()

    if st.session_state.result_text and not _result_task_active:
        # ── 修改摘要儀表板 ──────────────────────────────────
        _rt = st.session_state.result_text
        _old_snap = st.session_state.get("current_old_text_snapshot", "")
        _pending_q  = _rt.count("❓")
        _pending_w  = _rt.count("⚠️")
        _new_items  = _rt.count("▸")
        _del_items  = _rt.count("✕")
        _dash1, _dash2, _dash3, _dash4, _dash5 = st.columns(5)
        _dash1.metric("📝 原文字元", f"{len(_old_snap):,}" if _old_snap else "—")
        _dash2.metric("✨ 改寫字元", f"{len(_rt):,}")
        _dash3.metric("❓ 待確認項目", _pending_q)
        _dash4.metric("⚠️ 建議評估", _pending_w)
        _dash5.metric("▸ 新增 / ✕ 調整", f"{_new_items} / {_del_items}")
        st.markdown("---")

    if run_btn:
        if not api_key:
            if rewrite_progress_bar:
                rewrite_progress_bar.empty()
            if rewrite_status_box:
                rewrite_status_box.error("⚠️ 請先在左側輸入 Gemini API Key！")
            else:
                st.error("⚠️ 請先在左側輸入 Gemini API Key！")
        elif not old_text or not old_text.strip():
            if rewrite_progress_bar:
                rewrite_progress_bar.empty()
            if st.session_state.pdf_bytes_cache:
                _msg = "⚠️ 請先到「① 輸入」分頁點擊「📥 擷取選定頁面內容」，再執行改寫！"
            else:
                _msg = "⚠️ 請先於「① 輸入」分頁提供待改寫條文。"
            if rewrite_status_box:
                rewrite_status_box.error(_msg)
            else:
                st.error(_msg)
        else:
            st.session_state.running = True
            hint_section = (
                f"\n\n【使用者額外提示語（請優先遵照執行）】\n{user_hint.strip()}"
                if user_hint and user_hint.strip() else ""
            )
            protected_terms_input = st.session_state.get("protected_terms_area", "")
            protected_section = ""
            if protected_terms_input and protected_terms_input.strip():
                terms_list = [t.strip() for t in protected_terms_input.strip().splitlines() if t.strip()]
                if terms_list:
                    protected_section = (
                        "\n\n【禁止更改的專業術語（原文照錄，不得翻譯或替換）】\n"
                        + "\n".join(f"・{t}" for t in terms_list)
                    )
            kb_section = (
                f"\n\n【精進文件內容（請優先採用以下參數）】\n{kb_text}"
                if kb_text else "\n\n（本次未上傳精進文件，請依中運量通用規範改寫）"
            )
            prompt = f"""你是一位具備捷運機電系統工程與合約撰寫背景的資深專家。

【任務】
將下方『預修改之PTS（特別技術規範）』改寫為『未來線中運量系統』版本。

【改寫原則——請嚴格遵守，違反任何一條均為失敗】
1. 【範圍限制】你的改寫範圍僅限於「待改寫條文」區塊中所提供的文字，絕對不得自行補充、延伸或改寫未提供的章節內容。若待改寫條文的開頭有前文接續文字（如「（接續前頁）」），請直接從第一個完整條款開始改寫。
2. 嚴格保留原有章節編號與排版格式。
3. 【禁止省略】原文每一條、每一款、每一數值（溫度、面積、載重、時間等）都必須出現在改寫結果中，不得以「如原文」、「同上」或省略號代替。
4. 【表格必須完整保留】若原文有表格，改寫後必須以 Markdown 表格輸出，且嚴守以下格式規定：
   - 表格標題行（| 欄1 | 欄2 | 欄3 |）與分隔線（| --- | --- | --- |）必須完整輸出
   - 每一個測試項目（如 A.車廂地板防火測試、B.車間走道防火測試）都必須是表格中獨立的一列（row），不得拆到表格外
   - 說明欄中若有多個子條件（1. 2. 3.…），全部寫在同一格內，子條件之間用「；」分隔，不得換行成表格外的條列
   - 禁止在表格行之間插入任何非表格文字或空行
5. 優先採用精進文件中的最新中運量規格參數；若未提及，則保留原始數值，並將該條款編號與說明列入「=== 二、待確認事項 ===」，正文條文本體內不得出現〔待確認〕等標記。
6. 不適用未來線中運量的設備或條款，保留完整條文內容，並將該條款編號與不適用原因列入「=== 二、待確認事項 ===」的「建議處理」小節，正文條文本體內不得出現【建議刪除，或由機設處重新評估】等標記。
7. 每一條款之間必須空一行，保持段落清晰易讀；禁止將所有條文連成一段文字輸出。
8. 【禁止使用 HTML 標籤】輸出中禁止使用 <br>、<br/>、<p> 等任何 HTML 標籤，段落換行只能用空行（換兩次 Enter）表示。
9. 請勿輸出任何問候語，直接輸出以下四個區塊。{hint_section}{protected_section}{kb_section}

【待改寫條文】
{old_text}

---
請依以下格式輸出，每個區塊用 === 分隔：

=== 一、改寫後條文 ===
（完整改寫版條文，保留編號格式，各條款間空一行，表格用 Markdown 呈現）
（⚠️ 本區塊內容必須完全乾淨，禁止出現〔待確認〕、【建議刪除】等任何標記符號）

=== 二、待確認事項與建議處理項目 ===
【❓ 待確認項目】
（每項以「❓」開頭，格式：條款編號 → 待確認說明，各項間空一行）

【⚠️ 建議刪除或由機設處重新評估】
（每項以「⚠️」開頭，格式：條款編號 → 不適用原因，各項間空一行）

=== 三、新增內容摘要 ===
（每項以「▸」開頭，各項間空一行）

=== 四、刪除 / 調整內容摘要 ===
（每項以「✕」開頭，各項間空一行）

=== 五、專業意見與注意事項 ===
（每項以「💡」開頭，各項間空一行）
"""
            progress_bar = result_progress_bar or rewrite_progress_bar or st.progress(0, text="⏳ AI 改寫中，請稍候…")
            success, full_text = generate_with_progress(prompt, selected_model, api_key, progress_bar, "AI 改寫")
            if success:
                st.session_state.result_text = full_text
                chapter_label = st.session_state.get("chapter_id") or "手動輸入"
                _tw_now = datetime.datetime.utcnow() + datetime.timedelta(hours=8)
                _new_entry = {
                    "ts":       _tw_now.strftime("%m/%d %H:%M"),
                    "label":    chapter_label,
                    "text":     full_text,
                    "old_text": old_text,
                }
                # 重新賦值（非 in-place insert），確保 Streamlit session_state 偵測到變更
                st.session_state.rewrite_history = (
                    [_new_entry] + list(st.session_state.rewrite_history)
                )[:10]
                st.session_state.current_old_text_snapshot = old_text
                # 清除歷史選單 session state，確保下次顯示最新版
                st.session_state.pop("history_select", None)
            st.session_state.running = False
            if success:
                switch_to_result_page("✅ AI 改寫完成！已切換至「② 改寫結果」查看", "🎉")
            st.rerun()

    if _pending_refine_prompt:
        st.session_state.running = True
        progress_bar = result_progress_bar or st.progress(0, text="⏳ 二次精修中，請稍候…")
        success, full_text = generate_with_progress(_pending_refine_prompt, selected_model, api_key, progress_bar, "二次精修")
        if success:
            st.session_state.result_text = full_text
            chapter_label = st.session_state.get("chapter_id") or "手動輸入"
            _tw_now2 = datetime.datetime.utcnow() + datetime.timedelta(hours=8)
            _new_refine_entry = {
                "ts":       _tw_now2.strftime("%m/%d %H:%M"),
                "label":    f"{chapter_label}（精修）",
                "text":     full_text,
                "old_text": st.session_state.get("current_old_text_snapshot", ""),
            }
            st.session_state.rewrite_history = (
                [_new_refine_entry] + list(st.session_state.rewrite_history)
            )[:10]
            st.session_state.pop("history_select", None)
        st.session_state.pop("_pending_refine_prompt", None)
        st.session_state.running = False
        if success:
            switch_to_result_page("✅ 精修完成！已更新「② 改寫結果」。", "🔄")
        st.rerun()

    if st.session_state.result_text and not _result_task_active:
        # ── 版本 + 模式 控制列（分欄擺放，視覺更整齊） ──
        ctrl1, ctrl2 = st.columns([3, 2])
        history = st.session_state.rewrite_history
        with ctrl1:
            if len(history) > 1:
                history_labels = [f"第{i+1}版　{h['ts']}　{h['label']}" for i, h in enumerate(history)]
                sel_label = st.selectbox("📋 歷史版本", history_labels, index=0, key="history_select")
                sel_idx = history_labels.index(sel_label)
                display_result = history[sel_idx]["text"]
                display_old    = history[sel_idx]["old_text"]
            else:
                display_result = st.session_state.result_text
                display_old    = st.session_state.get("current_old_text_snapshot", "")
        with ctrl2:
            st.write("")
            compare_mode = st.toggle("📊 新舊對照模式", key="compare_toggle")

        if compare_mode:
            _secs = display_result.split("===")
            _parsed_for_compare = {}
            for _i in range(1, len(_secs) - 1, 2):
                _parsed_for_compare[_secs[_i].strip()] = (
                    _secs[_i + 1].strip() if _i + 1 < len(_secs) else "")
            section_one_text = next(
                (c for t, c in _parsed_for_compare.items() if t.startswith("一")), display_result)
            left_col, right_col = st.columns(2)
            with left_col:
                st.caption("📄 原始條文")
                st.text_area("原始", value=display_old or "（無原始條文快照）",
                    height=600, disabled=True, label_visibility="collapsed", key="orig_view")
            with right_col:
                st.caption("✨ 改寫後條文（一、）")
                st.text_area("改寫", value=section_one_text,
                    height=600, disabled=True, label_visibility="collapsed", key="rewrite_view")
        else:
            result_text = display_result
            sections = result_text.split("===")
            parsed = {}
            for i in range(1, len(sections) - 1, 2):
                title = sections[i].strip()
                content = sections[i + 1].strip() if i + 1 < len(sections) else ""
                parsed[title] = content
            emoji_map = {"一": "📄", "二": "⚠️", "三": "🆕", "四": "🗑️", "五": "💡"}

            def parse_md_table_to_html(md_lines):
                html = ['<table border="1" style="border-collapse:collapse;width:100%;font-size:0.9em;">']
                is_header = True
                for row in md_lines:
                    if re.match(r"^\|[-:\s|]+\|$", row.strip()):
                        is_header = False
                        continue
                    cells = [c.strip() for c in row.strip().strip("|").split("|")]
                    tag = "th" if is_header else "td"
                    style = "padding:6px 10px;vertical-align:top;border:1px solid #ccc;"
                    if is_header:
                        style += "background:#f0f0f0;font-weight:bold;"
                    row_html = "<tr>" + "".join(
                        f'<{tag} style="{style}">{format_cell(c)}</{tag}>' for c in cells) + "</tr>"
                    html.append(row_html)
                html.append("</table>")
                return "\n".join(html)

            def format_cell(text):
                import html as html_mod
                if text is None:
                    return ""
                text = str(text)
                text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
                text = text.replace("\r\n", "\n").replace("\r", "\n")
                text = re.sub(r"\s*；\s*", "；", text)
                item_pat = re.compile(r"(?<![\d.])(\d{1,2})[.、]\s*(?!\d)")
                matches = list(item_pat.finditer(text))
                if matches:
                    preamble = text[:matches[0].start()].strip("；; \n\t")
                    items = []
                    for idx, match in enumerate(matches):
                        item_start = match.end()
                        item_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
                        item_text = text[item_start:item_end].strip("；; \n\t")
                        if item_text:
                            items.append(item_text)
                    if items:
                        first_number = matches[0].group(1)
                        prefix = (f"<p style='margin:0 0 4px 0;'>{html_mod.escape(preamble)}</p>"
                                  if preamble else "")
                        li_items = "".join(
                            f"<li>{html_mod.escape(item).replace(chr(10), '<br>')}</li>" for item in items)
                        return (f"{prefix}<ol start='{first_number}' style='margin:0;padding-left:1.4em;'>"
                                f"{li_items}</ol>")
                return html_mod.escape(text).replace("\n", "<br>")

            def render_section(title, content):
                lines = content.split("\n")
                output_parts = []
                i = 0
                while i < len(lines):
                    stripped = lines[i].strip()
                    if stripped.startswith("|"):
                        tbl_rows = []
                        while i < len(lines) and (lines[i].strip().startswith("|") or
                              (not lines[i].strip() and i+1 < len(lines) and lines[i+1].strip().startswith("|"))):
                            if lines[i].strip():
                                tbl_rows.append(lines[i].strip())
                            i += 1
                        output_parts.append(("table", tbl_rows))
                    else:
                        md_lines = []
                        while i < len(lines) and not lines[i].strip().startswith("|"):
                            md_lines.append(lines[i])
                            i += 1
                        output_parts.append(("md", md_lines))
                final_html_parts = []
                for part_type, part_data in output_parts:
                    if part_type == "table":
                        final_html_parts.append(parse_md_table_to_html(part_data))
                    else:
                        md_text = ""
                        for ln in part_data:
                            s = ln.strip()
                            md_text += ln + "\n"
                            if s and not s.startswith("---") and not s.startswith("|"):
                                md_text += "\n"
                        if md_text.strip():
                            final_html_parts.append(f"MARKDOWN_BLOCK:{md_text}")
                for part in final_html_parts:
                    if part.startswith("MARKDOWN_BLOCK:"):
                        st.markdown(part[len("MARKDOWN_BLOCK:"):])
                    else:
                        st.markdown(part, unsafe_allow_html=True)

            if parsed:
                # ── 結果分欄：主條文（一）佔大欄，其餘摘要收進右側 ──
                main_title = next((t for t in parsed if t.startswith("一")), None)
                other_titles = [t for t in parsed if t != main_title]

                res_main, res_side = st.columns([3, 2], gap="large")
                with res_main:
                    if main_title:
                        st.markdown(f"#### {emoji_map.get('一','📄')} {main_title}")
                        render_section(main_title, parsed[main_title])
                with res_side:
                    st.markdown("#### 📑 摘要與待確認")
                    for title in other_titles:
                        emoji = emoji_map.get(title[0] if title else "", "📌")
                        with st.expander(f"{emoji} {title}", expanded=title.startswith("二")):
                            render_section(title, parsed[title])
            else:
                st.markdown(result_text)

        st.success("✅ 改寫完成！")
        dl1, dl2, dl3 = st.columns(3)
        with dl1:
            st.download_button("⬇️ 下載 Word 檔（.docx）",
                data=result_to_docx(display_result),
                file_name="未來線中運量_改寫結果.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        with dl2:
            st.download_button("⬇️ 下載純文字（.txt）",
                data=display_result.encode("utf-8"),
                file_name="未來線中運量_改寫結果.txt",
                mime="text/plain", use_container_width=True)
        with dl3:
            _html_content = f"""<!DOCTYPE html>
<html lang="zh-TW"><head><meta charset="UTF-8">
<title>未來線中運量_改寫結果</title>
<style>
  body{{font-family:'標楷體','DFKai-SB',serif;font-size:12pt;margin:2cm;line-height:1.8;}}
  h1{{font-size:16pt;color:#1F497D;border-bottom:2px solid #1F497D;padding-bottom:6px;}}
  h2{{font-size:13pt;color:#1F497D;}}
  table{{border-collapse:collapse;width:100%;margin:12px 0;}}
  th,td{{border:1px solid #999;padding:6px 10px;vertical-align:top;}}
  th{{background:#e8edf4;font-weight:bold;}}
  p{{margin:4px 0;}}
  @media print{{body{{margin:1.5cm;}} button{{display:none;}}}}
</style></head><body>
<h1>未來線中運量機電特別技術規範 — 智慧改寫結果</h1>
<pre style="white-space:pre-wrap;font-family:inherit;">{display_result}</pre>
</body></html>"""
            st.download_button("⬇️ 下載 HTML（瀏覽器可列印為 PDF）",
                data=_html_content.encode("utf-8"),
                file_name="未來線中運量_改寫結果.html",
                mime="text/html", use_container_width=True)
    elif not _result_task_active:
        st.info("尚未產出結果。請至「① 輸入」分頁提供條文並執行改寫。")

# ══════════════════════════════════════════════════════════
# ③ 二次精修分頁
# ══════════════════════════════════════════════════════════
with tab_refine:
    st.subheader("🔄 二次精修")
    if not st.session_state.result_text:
        st.info("尚無可精修的草稿，請先於「② 改寫結果」完成一次改寫。")
    else:
        st.caption("針對目前改寫結果輸入補充指令，AI 將依指令進行二次精修（結果同步更新至「② 改寫結果」）")
        result_text = st.session_state.result_text
        with st.expander("📄 目前草稿預覽（一、改寫後條文）", expanded=False):
            _secs = result_text.split("===")
            _p = {}
            for _i in range(1, len(_secs) - 1, 2):
                _p[_secs[_i].strip()] = _secs[_i + 1].strip() if _i + 1 < len(_secs) else ""
            _one = next((c for t, c in _p.items() if t.startswith("一")), result_text)
            st.text(_one[:2000] + ("..." if len(_one) > 2000 else ""))

        post_hint = st.text_area("輸入修改指令", height=140,
            placeholder="例如：\n・請將第 3.1.2 條的電壓改為 750V DC\n・把第 3.2 條表格的測試時間全部改為 30 分鐘\n・幫我刪除第 3.3 條整條",
            label_visibility="collapsed", key="post_hint_area")
        post_btn = st.button("🔄 依指令修改改寫結果", use_container_width=True,
            type="primary", disabled=st.session_state.running)

        if post_btn:
            if not post_hint or not post_hint.strip():
                st.warning("⚠️ 請先輸入修改指令！")
            elif not api_key:
                st.error("⚠️ 請先在左側輸入 Gemini API Key！")
            else:
                refine_prompt = f"""你是一位具備捷運機電系統工程與合約撰寫背景的資深專家。

【任務】
以下是一份已改寫完成的技術規範草稿，請依照「修改指令」對其進行調整，並完整保留未被指令影響的部分不變。

【修改指令】
{post_hint.strip()}

【現有改寫草稿】
{result_text}

【輸出要求】
- 請直接輸出修改後的完整草稿，格式與分區（=== 一、... === 二、... 等）保持不變
- 不要輸出任何問候語或說明，直接輸出修改後的草稿
- 一、改寫後條文 區塊內容必須乾淨，不得出現〔待確認〕、【建議刪除】等標記
"""
                st.session_state["_pending_refine_prompt"] = refine_prompt
                switch_to_result_page("⏳ 已切換至「② 改寫結果」執行二次精修。", "🔄")
                st.rerun()

